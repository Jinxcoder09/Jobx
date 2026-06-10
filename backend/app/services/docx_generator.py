import io
from typing import Any, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def clean_url(url: str, prefix: str = "https://") -> str:
    url = url.strip()
    if not url:
        return ""
    if url.startswith("http://") or url.startswith("https://") or url.startswith("mailto:"):
        return url
    if prefix == "mailto:" and "@" in url:
        return f"mailto:{url}"
    return f"{prefix}{url}"

def get_rgb_color(hex_str: str, default="000000") -> RGBColor:
    hex_str = hex_str.replace("#", "").strip()
    if len(hex_str) != 6:
        hex_str = hex_str.ljust(6, "0")[:6]
    try:
        return RGBColor.from_string(hex_str)
    except:
        return RGBColor.from_string(default)

def blend_with_white(hex_color: str, alpha=0.05) -> str:
    hex_color = hex_color.replace("#", "")
    if len(hex_color) != 6:
        return "F8FAFC"
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        
        r_new = int(r * alpha + 255 * (1 - alpha))
        g_new = int(g * alpha + 255 * (1 - alpha))
        b_new = int(b * alpha + 255 * (1 - alpha))
        
        return f"{r_new:02X}{g_new:02X}{b_new:02X}"
    except:
        return "F8FAFC"

def add_hyperlink(paragraph, url: str, text: str, color_hex: str = "2563eb", underline: bool = True):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = parse_xml(
        f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:id="{r_id}"/>'
    )
    new_run = parse_xml(f'<w:r {nsdecls("w")}/>')
    text_node = parse_xml(f'<w:t {nsdecls("w")}>{text}</w:t>')
    new_run.append(text_node)
    
    rPr = new_run.get_or_add_rPr()
    if color_hex:
        c = parse_xml(f'<w:color {nsdecls("w")} w:val="{color_hex.replace("#", "")}"/>')
        rPr.append(c)
    if underline:
        u = parse_xml(f'<w:u {nsdecls("w")} w:val="single"/>')
        rPr.append(u)
        
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def set_cell_background(cell, hex_color: str):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color.replace("#", "")}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def make_table_invisible(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def add_p_bottom_border(paragraph, color_hex="CCCCCC", size=12):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{size}" w:space="4" w:color="{color_hex.replace("#", "")}"/></w:pBdr>')
    pPr.append(pBdr)

class DocxResumeBuilder:
    def __init__(self, resume_data: dict):
        self.data = resume_data.get("data", {})
        self.theme = resume_data.get("theme", {})
        self.template_id = resume_data.get("templateId", "modern")
        
        self.font_family = self.theme.get("fontFamily", "Inter")
        self.font_size = self.theme.get("fontSize", 11)
        self.line_spacing = self.theme.get("lineSpacing", 1.4)
        self.section_spacing = self.theme.get("sectionSpacing", 16)
        
        self.primary_hex = self.theme.get("primaryColor", "#0f172a")
        self.secondary_hex = self.theme.get("secondaryColor", "#475569")
        self.accent_hex = self.theme.get("accentColor", "#2563eb")
        
        self.primary_color = get_rgb_color(self.primary_hex)
        self.secondary_color = get_rgb_color(self.secondary_hex)
        self.accent_color = get_rgb_color(self.accent_hex)
        
        self.layout_type = self.theme.get("layout", "single")

    def apply_run_font(self, run, size_offset=0, bold=False, italic=False, color_rgb=None):
        run.font.name = self.font_family
        run.font.size = Pt(self.font_size + size_offset)
        run.bold = bold
        run.italic = italic
        if color_rgb:
            run.font.color.rgb = color_rgb

    def add_styled_paragraph(self, container, text="", space_before=0, space_after=3, alignment=WD_ALIGN_PARAGRAPH.LEFT) -> Any:
        p = container.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = self.line_spacing
        if text:
            run = p.add_run(text)
            self.apply_run_font(run, color_rgb=self.secondary_color)
        return p

    def add_section_heading(self, container, title: str):
        t_id = self.template_id
        heading_text = title.upper() if t_id in ["elegant", "classic", "academic", "executive"] else title
        if t_id == "technical":
            heading_text = f"# {heading_text}"
            
        p = container.add_paragraph()
        p.paragraph_format.space_before = Pt(self.section_spacing)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(heading_text)
        self.apply_run_font(run, size_offset=1, bold=True, color_rgb=self.accent_color)
        
        add_p_bottom_border(p, color_hex=self.accent_hex, size=8)

    def build_header(self, container):
        p_info = self.data.get("personal", {})
        name = p_info.get("fullName", "Your Name")
        title = p_info.get("title", "")
        
        meta_parts = []
        if p_info.get("email"):
            meta_parts.append({"text": p_info["email"], "url": clean_url(p_info["email"], "mailto:"), "type": "link"})
        if p_info.get("phone"):
            meta_parts.append({"text": p_info["phone"], "type": "text"})
        if p_info.get("location"):
            meta_parts.append({"text": p_info["location"], "type": "text"})
        if p_info.get("website"):
            meta_parts.append({"text": p_info["website"], "url": clean_url(p_info["website"]), "type": "link"})
        if p_info.get("linkedin"):
            meta_parts.append({"text": "LinkedIn", "url": clean_url(p_info["linkedin"], "https://linkedin.com/in/"), "type": "link"})
        if p_info.get("github"):
            meta_parts.append({"text": "GitHub", "url": clean_url(p_info["github"], "https://github.com/"), "type": "link"})

        t_id = self.template_id

        # Centered Headers
        if t_id in ["elegant", "academic", "classic"]:
            name_p = self.add_styled_paragraph(container, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            run = name_p.add_run(name.upper())
            self.apply_run_font(run, size_offset=7, bold=True, color_rgb=self.primary_color)
            
            if title:
                title_p = self.add_styled_paragraph(container, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                run = title_p.add_run(title)
                self.apply_run_font(run, size_offset=1, italic=True, color_rgb=self.secondary_color)
                
            meta_p = self.add_styled_paragraph(container, space_before=4, space_after=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_p_bottom_border(meta_p, color_hex=self.accent_hex, size=6)
            
            first = True
            for part in meta_parts:
                if not first:
                    meta_p.add_run("   ·   ")
                first = False
                if part["type"] == "link":
                    add_hyperlink(meta_p, part["url"], part["text"], color_hex=self.accent_hex)
                else:
                    run = meta_p.add_run(part["text"])
                    self.apply_run_font(run, size_offset=-2, color_rgb=self.secondary_color)
            return

        # Executive Header
        if t_id == "executive":
            name_p = self.add_styled_paragraph(container)
            run = name_p.add_run(name)
            self.apply_run_font(run, size_offset=8, bold=True, color_rgb=self.primary_color)
            
            if title:
                title_p = self.add_styled_paragraph(container, space_after=4)
                run = title_p.add_run(title.upper())
                self.apply_run_font(run, size_offset=-1, bold=True, color_rgb=self.accent_color)
                
            meta_p = self.add_styled_paragraph(container, space_after=10)
            first = True
            for part in meta_parts:
                if not first:
                    meta_p.add_run("   |   ")
                first = False
                if part["type"] == "link":
                    add_hyperlink(meta_p, part["url"], part["text"], color_hex=self.accent_hex)
                else:
                    run = meta_p.add_run(part["text"])
                    self.apply_run_font(run, size_offset=-1, color_rgb=self.secondary_color)
            return

        # Bold Header
        if t_id == "bold":
            name_p = self.add_styled_paragraph(container)
            run = name_p.add_run(name)
            self.apply_run_font(run, size_offset=12, bold=True, color_rgb=self.primary_color)
            
            if title:
                title_p = self.add_styled_paragraph(container, space_after=4)
                run = title_p.add_run(title)
                self.apply_run_font(run, size_offset=1, bold=True, color_rgb=self.accent_color)
                
            meta_p = self.add_styled_paragraph(container, space_after=12)
            first = True
            for part in meta_parts:
                if not first:
                    meta_p.add_run("   ·   ")
                first = False
                if part["type"] == "link":
                    add_hyperlink(meta_p, part["url"], part["text"], color_hex=self.accent_hex)
                else:
                    run = meta_p.add_run(part["text"])
                    self.apply_run_font(run, size_offset=-1, color_rgb=self.secondary_color)
            return

        # Technical Header
        if t_id == "technical":
            name_p = self.add_styled_paragraph(container)
            run = name_p.add_run(name)
            self.apply_run_font(run, size_offset=5, bold=True, color_rgb=self.primary_color)
            
            if title:
                title_p = self.add_styled_paragraph(container, space_after=4)
                run = title_p.add_run(f"// {title}")
                self.apply_run_font(run, color_rgb=self.accent_color)
                
            meta_p = self.add_styled_paragraph(container, space_after=10)
            first = True
            for part in meta_parts:
                if not first:
                    meta_p.add_run("   ·   ")
                first = False
                if part["type"] == "link":
                    add_hyperlink(meta_p, part["url"], part["text"], color_hex=self.accent_hex)
                else:
                    run = meta_p.add_run(part["text"])
                    self.apply_run_font(run, color_rgb=self.secondary_color)
            return

        # Default / Modern / Creative Header
        name_p = self.add_styled_paragraph(container)
        run = name_p.add_run(name)
        self.apply_run_font(run, size_offset=6, bold=True, color_rgb=self.primary_color)
        
        if title:
            title_p = self.add_styled_paragraph(container, space_after=4)
            run = title_p.add_run(title)
            self.apply_run_font(run, bold=True, color_rgb=self.accent_color)
            
        meta_p = self.add_styled_paragraph(container, space_after=10)
        first = True
        for part in meta_parts:
            if not first:
                meta_p.add_run("   ·   ")
            first = False
            if part["type"] == "link":
                add_hyperlink(meta_p, part["url"], part["text"], color_hex=self.accent_hex)
            else:
                run = meta_p.add_run(part["text"])
                self.apply_run_font(run, color_rgb=self.secondary_color)

    def build_summary(self, container):
        summary_text = self.data.get("summary", "")
        if not summary_text:
            return
        self.add_section_heading(container, "Summary")
        self.add_styled_paragraph(container, summary_text)

    def build_experience(self, container):
        exp = self.data.get("experience", [])
        if not exp:
            return
        self.add_section_heading(container, "Experience")
        for item in exp:
            company = item.get("company", "")
            role = item.get("role", "")
            loc = item.get("location", "")
            start = item.get("startDate", "")
            end = "Present" if item.get("current") else item.get("endDate", "")
            date_str = " – ".join(filter(None, [start, end]))
            
            p = self.add_styled_paragraph(container, space_before=4, space_after=1)
            role_run = p.add_run(role)
            self.apply_run_font(role_run, bold=True, color_rgb=self.primary_color)
            
            if company:
                p.add_run("  —  ")
                comp_run = p.add_run(company)
                self.apply_run_font(comp_run, color_rgb=self.secondary_color)
                
            meta_p = self.add_styled_paragraph(container, space_before=0, space_after=2)
            meta_parts = []
            if date_str:
                meta_parts.append(date_str)
            if loc:
                meta_parts.append(loc)
            meta_run = meta_p.add_run("  ·  ".join(meta_parts))
            self.apply_run_font(meta_run, size_offset=-1, italic=True, color_rgb=self.secondary_color)
            
            for bullet in item.get("bullets", []):
                if bullet.strip():
                    bp = container.add_paragraph(style='List Bullet')
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after = Pt(2)
                    bp.paragraph_format.line_spacing = self.line_spacing
                    bp_run = bp.add_run(bullet.strip())
                    self.apply_run_font(bp_run, color_rgb=self.secondary_color)

    def build_education(self, container):
        edu = self.data.get("education", [])
        if not edu:
            return
        self.add_section_heading(container, "Education")
        for item in edu:
            school = item.get("school", "")
            degree = item.get("degree", "")
            field = item.get("field", "")
            loc = item.get("location", "")
            start = item.get("startDate", "")
            end = item.get("endDate", "")
            date_str = " – ".join(filter(None, [start, end]))
            gpa = item.get("gpa", "")
            
            degree_field = ", ".join(filter(None, [degree, field]))
            header_left = degree_field if degree_field else school
            
            p = self.add_styled_paragraph(container, space_before=4, space_after=1)
            deg_run = p.add_run(header_left)
            self.apply_run_font(deg_run, bold=True, color_rgb=self.primary_color)
            
            meta_p = self.add_styled_paragraph(container, space_before=0, space_after=2)
            meta_parts = []
            if degree_field and school:
                meta_parts.append(school)
            if loc:
                meta_parts.append(loc)
            if date_str:
                meta_parts.append(date_str)
            if gpa:
                meta_parts.append(f"GPA: {gpa}")
                
            meta_run = meta_p.add_run("  ·  ".join(meta_parts))
            self.apply_run_font(meta_run, size_offset=-1, italic=True, color_rgb=self.secondary_color)
            
            desc = item.get("description", "")
            if desc:
                self.add_styled_paragraph(container, desc)

    def build_projects(self, container):
        proj = self.data.get("projects", [])
        if not proj:
            return
        self.add_section_heading(container, "Projects")
        for item in proj:
            name = item.get("name", "")
            link = item.get("link", "")
            desc = item.get("description", "")
            tech = item.get("technologies", [])
            
            p = self.add_styled_paragraph(container, space_before=4, space_after=2)
            name_run = p.add_run(name)
            self.apply_run_font(name_run, bold=True, color_rgb=self.primary_color)
            
            if link:
                p.add_run("  ·  ")
                add_hyperlink(p, clean_url(link), link, color_hex=self.accent_hex)
                
            if desc:
                self.add_styled_paragraph(container, desc)
                
            for bullet in item.get("bullets", []):
                if bullet.strip():
                    bp = container.add_paragraph(style='List Bullet')
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after = Pt(2)
                    bp.paragraph_format.line_spacing = self.line_spacing
                    bp_run = bp.add_run(bullet.strip())
                    self.apply_run_font(bp_run, color_rgb=self.secondary_color)
                    
            if tech:
                tech_p = self.add_styled_paragraph(container, space_before=2)
                tech_run = tech_p.add_run(f"Technologies: {', '.join(tech)}")
                self.apply_run_font(tech_run, size_offset=-1, italic=True, color_rgb=self.secondary_color)

    def build_skills(self, container):
        skills = self.data.get("skills", [])
        if not skills:
            return
        self.add_section_heading(container, "Skills")
        for grp in skills:
            cat = grp.get("category", "Skills")
            items = grp.get("items", [])
            if items:
                p = self.add_styled_paragraph(container)
                cat_run = p.add_run(f"{cat}: ")
                self.apply_run_font(cat_run, bold=True, color_rgb=self.primary_color)
                items_run = p.add_run(", ".join(items))
                self.apply_run_font(items_run, color_rgb=self.secondary_color)

    def build_simple_list(self, container, heading: str, items: List[dict]):
        if not items:
            return
        self.add_section_heading(container, heading)
        for item in items:
            title = item.get("title", "")
            subtitle = item.get("subtitle", "")
            date = item.get("date", "")
            desc = item.get("description", "")
            
            p = self.add_styled_paragraph(container, space_before=4, space_after=1)
            t_run = p.add_run(title)
            self.apply_run_font(t_run, bold=True, color_rgb=self.primary_color)
            
            if subtitle:
                p.add_run("  —  ")
                st_run = p.add_run(subtitle)
                self.apply_run_font(st_run, color_rgb=self.secondary_color)
                
            if date:
                meta_p = self.add_styled_paragraph(container, space_before=0, space_after=2)
                meta_run = meta_p.add_run(date)
                self.apply_run_font(meta_run, size_offset=-1, italic=True, color_rgb=self.secondary_color)
                
            if desc:
                self.add_styled_paragraph(container, desc)

    def build_languages(self, container):
        langs = self.data.get("languages", [])
        if not langs:
            return
        self.add_section_heading(container, "Languages")
        parts = []
        for l in langs:
            name = l.get("name", "")
            lvl = l.get("level", "")
            parts.append(f"{name} ({lvl})" if lvl else name)
            
        p = self.add_styled_paragraph(container)
        run = p.add_run("  ·  ".join(parts))
        self.apply_run_font(run, color_rgb=self.secondary_color)

    def build_custom_sections(self, container):
        customs = self.data.get("custom", [])
        if not customs:
            return
        for cs in customs:
            title = cs.get("title", "Custom Section")
            items = cs.get("items", [])
            self.build_simple_list(container, title, items)

    def build_section(self, container, key: str):
        if key == "summary":
            self.build_summary(container)
        elif key == "experience":
            self.build_experience(container)
        elif key == "education":
            self.build_education(container)
        elif key == "projects":
            self.build_projects(container)
        elif key == "skills":
            self.build_skills(container)
        elif key == "certifications":
            self.build_simple_list(container, "Certifications", self.data.get("certifications", []))
        elif key == "achievements":
            self.build_simple_list(container, "Achievements", self.data.get("achievements", []))
        elif key == "languages":
            self.build_languages(container)

    def generate(self) -> bytes:
        doc = Document()
        
        # Configure page margins
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
            
        # Configure base styles
        style = doc.styles['Normal']
        font = style.font
        font.name = self.font_family
        font.size = Pt(self.font_size)
        
        order = self.data.get("sectionOrder", [
            "summary", "experience", "education", "projects",
            "skills", "certifications", "achievements", "languages"
        ])
        
        # 1. Header (Personal Info) is always top
        self.build_header(doc)
        
        # 2. Check layout styling
        is_two_col = (
            self.layout_type == "two-column" or
            self.template_id in ["creative", "two-column", "timeline"]
        )
        
        if is_two_col:
            # 2-Column layout: create a table with 1 row and 2 columns
            table = doc.add_table(rows=1, cols=2)
            make_table_invisible(table)
            
            is_creative = (self.template_id == "creative")
            col_widths = [Inches(2.4), Inches(4.7)] if is_creative else [Inches(4.7), Inches(2.4)]
            
            table.columns[0].width = col_widths[0]
            table.columns[1].width = col_widths[1]
            
            left_cell = table.cell(0, 0)
            right_cell = table.cell(0, 1)
            
            left_cell.width = col_widths[0]
            right_cell.width = col_widths[1]
            
            sidebar_keys = ["skills", "languages", "certifications", "achievements"]
            
            if is_creative:
                left_keys = [k for k in order if k in sidebar_keys]
                right_keys = [k for k in order if k not in sidebar_keys]
                
                bg_color = blend_with_white(self.accent_hex, alpha=0.06)
                set_cell_background(left_cell, bg_color)
            else:
                left_keys = [k for k in order if k not in sidebar_keys]
                right_keys = [k for k in order if k in sidebar_keys]
                
            for k in left_keys:
                self.build_section(left_cell, k)
            if not is_creative:
                self.build_custom_sections(left_cell)
                
            for k in right_keys:
                self.build_section(right_cell, k)
            if is_creative:
                self.build_custom_sections(right_cell)
        else:
            for key in order:
                self.build_section(doc, key)
            self.build_custom_sections(doc)
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
